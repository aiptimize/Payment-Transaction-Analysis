# import os
# import pandas as pd
# import pdfplumber
# from flask import Flask, render_template, request, redirect, url_for, send_file, session
# from werkzeug.utils import secure_filename
# from docx import Document
# from docx.shared import Pt
# from docx.shared import Inches
# import matplotlib.pyplot as plt
# import matplotlib.font_manager as fm

# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
# app.config['OUTPUT_FOLDER'] = 'output'
# app.config['SECRET_KEY'] = 'your-secret-key-here'
# app.config['CHART_FOLDER'] = 'charts'

# # 创建上传、输出和图表目录
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
# os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
# os.makedirs(app.config['CHART_FOLDER'], exist_ok=True)

# # 定义敏感词和行业关键词的默认值
# DEFAULT_SENSITIVE_WORDS = ["法院"]
# DEFAULT_INDUSTRY_KEYWORDS = ["水泥"]

# # 设置matplotlib字体
# plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统中实际字体修改
# plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# @app.route('/', methods=['GET', 'POST'])
# def upload_file():
#     if request.method == 'POST':
#         # 获取表单数据
#         pdf_file = request.files['pdf_file']
#         sensitive_words = request.form.get('sensitive_words', '').split(',')
#         industry_keywords = request.form.get('industry_keywords', '').split(',')

#         # 去除空白关键词
#         sensitive_words = [word.strip() for word in sensitive_words if word.strip()]
#         industry_keywords = [word.strip() for word in industry_keywords if word.strip()]

#         # 如果用户没有输入关键词，使用默认值
#         if not sensitive_words:
#             sensitive_words = DEFAULT_SENSITIVE_WORDS
#         if not industry_keywords:
#             industry_keywords = DEFAULT_INDUSTRY_KEYWORDS

#         # 保存上传的PDF文件
#         if pdf_file:
#             filename = secure_filename(pdf_file.filename)
#             input_pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             pdf_file.save(input_pdf_path)

#             # 处理PDF文件
#             output_filename = f"微信支付交易明细报告.docx"  # 修改为docx扩展名
#             output_doc_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

#             try:
#                 # 确保输出目录存在
#                 os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
                
#                 process_pdf(input_pdf_path, output_doc_path, sensitive_words, industry_keywords)
                
#                 # 检查文件是否成功生成
#                 if os.path.exists(output_doc_path):
#                     # 将文件名存储在session中，以便在另一个路由中使用
#                     session['doc_filename'] = output_filename
#                     return redirect(url_for('upload_file'))  # 重定向回上传页面
#                 else:
#                     raise Exception("DOCX文件未成功生成，但没有抛出具体错误")
                    
#             except Exception as e:
#                 error_msg = f"处理文件时出错: {str(e)}"
#                 print(f"错误详情: {str(e)}")  # 打印详细错误信息到控制台
#                 return render_template('upload.html', error=error_msg,
#                                        default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
#                                        default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
#                                        session=session)

#     # 渲染上传页面
#     return render_template('upload.html',
#                            default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
#                            default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
#                            session=session)

# @app.route('/download')
# def download_file():
#     if 'doc_filename' not in session:
#         return redirect(url_for('upload_file'))

#     doc_filename = session['doc_filename']
#     doc_path = os.path.join(app.config['OUTPUT_FOLDER'], doc_filename)

#     if os.path.exists(doc_path):
#         return send_file(doc_path, as_attachment=True)
#     else:
#         # 清理无效的session数据
#         session.pop('doc_filename', None)
#         return render_template('upload.html', error="生成的DOCX文件不存在，请重新上传和分析",
#                                default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
#                                default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
#                                session=session)

# def process_pdf(input_pdf_path, output_doc_path, sensitive_words, industry_keywords):
#     try:
#         # 验证输入文件存在
#         if not os.path.exists(input_pdf_path):
#             raise FileNotFoundError(f"输入的PDF文件不存在: {input_pdf_path}")

#         all_data = []
#         header = None

#         with pdfplumber.open(input_pdf_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 # 提取当前页表格
#                 table = page.extract_table()

#                 if table:
#                     if i == 0:  # 假设第一页包含表头
#                         header = table[0]
#                         all_data.extend(table[1:])  # 添加数据行
#                     else:
#                         # 跳过后续页的重复表头（根据实际情况调整）
#                         all_data.extend(table[1:] if table[0] == header else table)

#         # 生成DataFrame并保存
#         if all_data:
#             # 删除前3行数据
#             all_data = all_data[2:]

#             # 定义新标题
#             new_columns = [
#                 "交易单号",
#                 "交易时间",
#                 "交易类型",
#                 "收/支/其他",
#                 "交易方式",
#                 "金额(元)",
#                 "交易对方",
#                 "商户单号"
#             ]

#             df = pd.DataFrame(all_data, columns=new_columns)
#         else:
#             raise ValueError("未找到表格数据")

#         # 数据处理
#         df = df.apply(lambda x: x.map(lambda y: y.replace('\n', '') if isinstance(y, str) else y))

#         # 确保金额列是数值类型
#         df['金额(元)'] = pd.to_numeric(df['金额(元)'], errors='coerce')

#         # 创建Word文档对象
#         doc = Document()

#         def add_text_to_doc(text, font_size=10, bold=False):
#             paragraph = doc.add_paragraph()
#             run = paragraph.add_run(text)
#             run.font.size = Pt(font_size)
#             run.font.bold = bold

#         # 写入标题
#         add_text_to_doc("微信支付交易明细分析报告", font_size=16, bold=True)
#         add_text_to_doc("-" * 80)
#         doc.add_paragraph()  # 增加间距

#         # 添加目录
#         doc.add_heading('目录', level=1)
#         doc.add_paragraph('1. 敏感词消费记录\t\t\t\t页 1')
#         doc.add_paragraph('2. 行业消费记录\t\t\t\t页 2')
#         doc.add_paragraph('3. 区域消费金额统计\t\t\t\t页 3')
#         doc.add_paragraph('4. 交易金额前5的交易对方\t\t\t页 4')
#         doc.add_paragraph('5. 交易数量前5的交易对方\t\t\t页 5')
#         doc.add_page_break()


#         # 筛选包含敏感词的交易并创建副本
#         filtered_df = df[df["交易对方"].str.contains('|'.join(sensitive_words), na=False)].copy()

#         # 删除交易单号列和商户单号列
#         columns_to_drop = ['交易单号', '商户单号']
#         for col in columns_to_drop:
#             if col in filtered_df.columns:
#                 filtered_df = filtered_df.drop(columns=[col])

#         # 检查是否有符合条件的记录
#         if not filtered_df.empty:
#             # 按金额倒序排列
#             sorted_df = filtered_df.dropna(subset=["金额(元)"]).sort_values(by="金额(元)", ascending=False)

#             # 重置索引并打印结果
#             sorted_df.reset_index(drop=True, inplace=True)
#             doc.add_heading('1. 敏感词消费记录', level=1)
#             add_text_to_doc("敏感词消费记录（按金额降序排列）：", font_size=12, bold=True)
            
#             # 创建表格
#             table = doc.add_table(rows=1, cols=len(sorted_df.columns))
#             hdr_cells = table.rows[0].cells
            
#             # 添加表头
#             for i, col in enumerate(sorted_df.columns):
#                 hdr_cells[i].text = col
            
#             # 添加数据行
#             for _, row in sorted_df.head(5).iterrows():
#                 row_cells = table.add_row().cells
#                 for i, value in enumerate(row):
#                     row_cells[i].text = str(value)
#         else:
#             doc.add_heading('1. 敏感词消费记录', level=1)
#             add_text_to_doc("没有敏感词消费记录")

#         doc.add_page_break()

#         # 筛选包含行业关键词的交易并创建副本
#         filtered_df = df[df["交易对方"].str.contains('|'.join(industry_keywords), na=False, case=False)].copy()

#         # 删除交易单号列和商户单号列
#         columns_to_drop = ['交易单号', '商户单号']
#         for col in columns_to_drop:
#             if col in filtered_df.columns:
#                 filtered_df = filtered_df.drop(columns=[col])

#         # 检查是否有符合条件的记录
#         if not filtered_df.empty:
#             # 按金额倒序排列
#             sorted_df = filtered_df.dropna(subset=["金额(元)"]).sort_values(by="金额(元)", ascending=False)

#             # 重置索引并打印结果
#             sorted_df.reset_index(drop=True, inplace=True)
#             doc.add_heading('2. 行业消费记录', level=1)
#             add_text_to_doc("行业消费记录（按金额降序排列）：", font_size=12, bold=True)
            
#             # 创建表格
#             table = doc.add_table(rows=1, cols=len(sorted_df.columns))
#             hdr_cells = table.rows[0].cells
            
#             # 添加表头
#             for i, col in enumerate(sorted_df.columns):
#                 hdr_cells[i].text = col
            
#             # 添加数据行
#             for _, row in sorted_df.head(5).iterrows():
#                 row_cells = table.add_row().cells
#                 for i, value in enumerate(row):
#                     row_cells[i].text = str(value)
#         else:
#             doc.add_heading('2. 行业消费记录', level=1)
#             add_text_to_doc("没有行业消费记录")

#         doc.add_page_break()

#         # 定义区域关键词库
#         location_keywords = ["潢川县", "深圳市", "郑州市", "温县"]  # 可根据需要添加更多地区

#         # 尝试可能的列名变体
#         column_variants = ['交易对方', '对方', '收款方', '商户名称', 'counterpart', 'recipient']
#         target_column = next((col for col in column_variants if col in df.columns), None)

#         if target_column is None:
#             available_cols = ", ".join(df.columns.tolist())
#             raise KeyError(f"找不到交易对方列，可用的列有：{available_cols}")

#         # 筛选包含地理关键词的交易并创建副本
#         filtered_df = df[df[target_column].str.contains('|'.join(location_keywords), na=False, case=False, regex=True)].copy()

#         # 检查金额列名称
#         amount_variants = ['金额(元)', '金额', 'money', 'amount', '交易金额']
#         amount_column = next((col for col in amount_variants if col in df.columns), None)

#         if amount_column is None:
#             available_cols = ", ".join(df.columns.tolist())
#             raise KeyError(f"找不到金额列，可用的列有：{available_cols}")

#         # 按金额倒序排列
#         sorted_df = filtered_df.dropna(subset=[amount_column]).sort_values(by=amount_column, ascending=False)

#         # 设置显示选项
#         pd.set_option('display.max_rows', None)
#         pd.set_option('display.unicode.ambiguous_as_wide', True)
#         pd.set_option('display.unicode.east_asian_width', True)

#         # 提取区域信息并统计前5个区域
#         def extract_location(text, keywords):
#             for keyword in keywords:
#                 if keyword in str(text):
#                     return keyword
#             return "其他"

#         sorted_df['区域'] = sorted_df[target_column].apply(lambda x: extract_location(x, location_keywords))

#         # 按区域分组并计算总金额
#         region_stats = sorted_df.groupby('区域')[amount_column].sum().sort_values(ascending=False).head(5)

#         # 打印结果
#         doc.add_heading('3. 区域消费金额统计', level=1)
#         add_text_to_doc("区域消费金额统计：", font_size=12, bold=True)
        
#         # 创建表格
#         table = doc.add_table(rows=1, cols=2)
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = '区域'
#         hdr_cells[1].text = '金额(元)'
        
#         # 添加数据行
#         for region, amount in region_stats.items():
#             row_cells = table.add_row().cells
#             row_cells[0].text = region
#             row_cells[1].text = f"{amount:.2f}"

#         # 绘制区域消费金额统计柱状图
#         plt.figure(figsize=(10, 6))
#         plt.bar(region_stats.index, region_stats.values)
#         plt.xlabel('区域')
#         plt.ylabel('金额(元)')
#         plt.title('区域消费金额统计')
#         chart_path = os.path.join(app.config['CHART_FOLDER'], 'region_chart.png')
#         plt.savefig(chart_path)
#         plt.close()
#         doc.add_picture(chart_path, width=Inches(6))

#         doc.add_page_break()

#         # 按交易对方分组，计算交易金额总和，并筛选出交易金额前5的交易对方
#         # 确保在进行数值计算前金额列已经转换为数值类型
#         top_5_by_amount = df.groupby('交易对方')['金额(元)'].sum().abs().sort_values(ascending=False).head(5)

#         # 按交易对方分组，计算交易数量，并筛选出交易数量前5的交易对方
#         top_5_by_count = df.groupby('交易对方')['金额(元)'].count().sort_values(ascending=False).head(5)

#         doc.add_heading('4. 交易金额前5的交易对方', level=1)
#         add_text_to_doc('交易金额前5的交易对方：', font_size=12, bold=True)
        
#         # 创建表格
#         table = doc.add_table(rows=1, cols=2)
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = '交易对方'
#         hdr_cells[1].text = '金额(元)'
        
#         # 添加数据行
#         for entity, amount in top_5_by_amount.items():
#             row_cells = table.add_row().cells
#             row_cells[0].text = entity
#             row_cells[1].text = f"{amount:.2f}"

#         # 绘制交易金额前5的交易对方柱状图
#         plt.figure(figsize=(10, 6))
#         plt.bar(top_5_by_amount.index, top_5_by_amount.values)
#         plt.xlabel('交易对方')
#         plt.ylabel('金额(元)')
#         plt.title('交易金额前5的交易对方')
#         chart_path = os.path.join(app.config['CHART_FOLDER'], 'top_5_amount_chart.png')
#         plt.savefig(chart_path)
#         plt.close()
#         doc.add_picture(chart_path, width=Inches(6))

#         doc.add_page_break()

#         doc.add_heading('5. 交易数量前5的交易对方', level=1)
#         add_text_to_doc('交易数量前5的交易对方：', font_size=12, bold=True)
        
#         # 创建表格
#         table = doc.add_table(rows=1, cols=2)
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = '交易对方'
#         hdr_cells[1].text = '交易数量'
        
#         # 添加数据行
#         for entity, count in top_5_by_count.items():
#             row_cells = table.add_row().cells
#             row_cells[0].text = entity
#             row_cells[1].text = str(count)

#         # 绘制交易数量前5的交易对方柱状图
#         plt.figure(figsize=(10, 6))
#         plt.bar(top_5_by_count.index, top_5_by_count.values)
#         plt.xlabel('交易对方')
#         plt.ylabel('交易数量')
#         plt.title('交易数量前5的交易对方')
#         chart_path = os.path.join(app.config['CHART_FOLDER'], 'top_5_count_chart.png')
#         plt.savefig(chart_path)
#         plt.close()
#         doc.add_picture(chart_path, width=Inches(6))

#         # 保存Word文档
#         doc.save(output_doc_path)

#         # 验证文件是否成功保存
#         if not os.path.exists(output_doc_path):
#             raise Exception(f"DOCX保存失败，文件不存在: {output_doc_path}")

#     except Exception as e:
#         # 错误发生时尝试删除不完整的文件
#         if os.path.exists(output_doc_path):
#             try:
#                 os.remove(output_doc_path)
#             except:
#                 pass
#         raise  # 重新抛出异常，让上层处理

# if __name__ == '__main__':
#     app.run(debug=True)




import os
import pandas as pd
import pdfplumber
from flask import Flask, render_template, request, redirect, url_for, send_file, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['CHART_FOLDER'] = 'charts'

# 创建上传、输出和图表目录
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
os.makedirs(app.config['CHART_FOLDER'], exist_ok=True)

# 定义敏感词和行业关键词的默认值
DEFAULT_SENSITIVE_WORDS = ["法院"]
DEFAULT_INDUSTRY_KEYWORDS = ["水泥"]

# 设置matplotlib字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统中实际字体修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # 获取表单数据
        pdf_file = request.files['pdf_file']
        sensitive_words = request.form.get('sensitive_words', '').split(',')
        industry_keywords = request.form.get('industry_keywords', '').split(',')

        # 去除空白关键词
        sensitive_words = [word.strip() for word in sensitive_words if word.strip()]
        industry_keywords = [word.strip() for word in industry_keywords if word.strip()]

        # 如果用户没有输入关键词，使用默认值
        if not sensitive_words:
            sensitive_words = DEFAULT_SENSITIVE_WORDS
        if not industry_keywords:
            industry_keywords = DEFAULT_INDUSTRY_KEYWORDS

        # 保存上传的PDF文件
        if pdf_file:
            filename = secure_filename(pdf_file.filename)
            input_pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            pdf_file.save(input_pdf_path)

            # 处理PDF文件
            output_filename = f"微信支付交易明细报告.docx"  # 修改为docx扩展名
            output_doc_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            try:
                # 确保输出目录存在
                os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
                
                process_pdf(input_pdf_path, output_doc_path, sensitive_words, industry_keywords)
                
                # 检查文件是否成功生成
                if os.path.exists(output_doc_path):
                    # 将文件名存储在session中，以便在另一个路由中使用
                    session['doc_filename'] = output_filename
                    return redirect(url_for('upload_file'))  # 重定向回上传页面
                else:
                    raise Exception("DOCX文件未成功生成，但没有抛出具体错误")
                    
            except Exception as e:
                error_msg = f"处理文件时出错: {str(e)}"
                print(f"错误详情: {str(e)}")  # 打印详细错误信息到控制台
                return render_template('upload.html', error=error_msg,
                                       default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
                                       default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
                                       session=session)

    # 渲染上传页面
    return render_template('upload.html',
                           default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
                           default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
                           session=session)

@app.route('/download')
def download_file():
    if 'doc_filename' not in session:
        return redirect(url_for('upload_file'))

    doc_filename = session['doc_filename']
    doc_path = os.path.join(app.config['OUTPUT_FOLDER'], doc_filename)

    if os.path.exists(doc_path):
        return send_file(doc_path, as_attachment=True)
    else:
        # 清理无效的session数据
        session.pop('doc_filename', None)
        return render_template('upload.html', error="生成的DOCX文件不存在，请重新上传和分析",
                               default_sensitive_words=','.join(DEFAULT_SENSITIVE_WORDS),
                               default_industry_keywords=','.join(DEFAULT_INDUSTRY_KEYWORDS),
                               session=session)

def process_pdf(input_pdf_path, output_doc_path, sensitive_words, industry_keywords):
    try:
        # 验证输入文件存在
        if not os.path.exists(input_pdf_path):
            raise FileNotFoundError(f"输入的PDF文件不存在: {input_pdf_path}")

        all_data = []
        header = None

        with pdfplumber.open(input_pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # 提取当前页表格
                table = page.extract_table()

                if table:
                    if i == 0:  # 假设第一页包含表头
                        header = table[0]
                        all_data.extend(table[1:])  # 添加数据行
                    else:
                        # 跳过后续页的重复表头（根据实际情况调整）
                        all_data.extend(table[1:] if table[0] == header else table)

        # 生成DataFrame并保存
        if all_data:
            # 删除前3行数据
            all_data = all_data[2:]

            # 定义新标题
            new_columns = [
                "交易单号",
                "交易时间",
                "交易类型",
                "收/支/其他",
                "交易方式",
                "金额(元)",
                "交易对方",
                "商户单号"
            ]

            df = pd.DataFrame(all_data, columns=new_columns)
        else:
            raise ValueError("未找到表格数据")

        # 数据处理
        df = df.apply(lambda x: x.map(lambda y: y.replace('\n', '') if isinstance(y, str) else y))

        # 确保金额列是数值类型
        df['金额(元)'] = pd.to_numeric(df['金额(元)'], errors='coerce')

        # 创建Word文档对象
        doc = Document()

        def add_text_to_doc(text, font_size=10, bold=False):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = bold

        # 写入标题
        add_text_to_doc("微信支付交易明细分析报告", font_size=16, bold=True)
        add_text_to_doc("-" * 80)
        doc.add_paragraph()  # 增加间距

        # 添加目录
        doc.add_heading('目录', level=1)
        doc.add_paragraph('1. 敏感词消费记录\t\t\t\t页 1')
        doc.add_paragraph('2. 行业消费记录\t\t\t\t页 2')
        doc.add_paragraph('3. 区域消费金额统计\t\t\t\t页 3')
        doc.add_paragraph('4. 交易金额前5的交易对方\t\t\t页 4')
        doc.add_paragraph('5. 交易数量前5的交易对方\t\t\t页 5')
        doc.add_page_break()


        # 筛选包含敏感词的交易并创建副本
        filtered_df = df[df["交易对方"].str.contains('|'.join(sensitive_words), na=False)].copy()

        # 删除交易单号列和商户单号列
        columns_to_drop = ['交易单号', '商户单号']
        for col in columns_to_drop:
            if col in filtered_df.columns:
                filtered_df = filtered_df.drop(columns=[col])

        # 检查是否有符合条件的记录
        if not filtered_df.empty:
            # 按金额倒序排列
            sorted_df = filtered_df.dropna(subset=["金额(元)"]).sort_values(by="金额(元)", ascending=False)

            # 重置索引并打印结果
            sorted_df.reset_index(drop=True, inplace=True)
            doc.add_heading('1. 敏感词消费记录', level=1)
            add_text_to_doc("敏感词消费记录（按金额降序排列）：", font_size=12, bold=True)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=len(sorted_df.columns))
            hdr_cells = table.rows[0].cells
            
            # 添加表头
            for i, col in enumerate(sorted_df.columns):
                hdr_cells[i].text = col
            
            # 添加数据行
            for _, row in sorted_df.head(5).iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        else:
            doc.add_heading('1. 敏感词消费记录', level=1)
            add_text_to_doc("没有敏感词消费记录")

        doc.add_page_break()

        # 筛选包含行业关键词的交易并创建副本
        filtered_df = df[df["交易对方"].str.contains('|'.join(industry_keywords), na=False, case=False)].copy()

        # 删除交易单号列和商户单号列
        columns_to_drop = ['交易单号', '商户单号']
        for col in columns_to_drop:
            if col in filtered_df.columns:
                filtered_df = filtered_df.drop(columns=[col])

        # 检查是否有符合条件的记录
        if not filtered_df.empty:
            # 按金额倒序排列
            sorted_df = filtered_df.dropna(subset=["金额(元)"]).sort_values(by="金额(元)", ascending=False)

            # 重置索引并打印结果
            sorted_df.reset_index(drop=True, inplace=True)
            doc.add_heading('2. 行业消费记录', level=1)
            add_text_to_doc("行业消费记录（按金额降序排列）：", font_size=12, bold=True)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=len(sorted_df.columns))
            hdr_cells = table.rows[0].cells
            
            # 添加表头
            for i, col in enumerate(sorted_df.columns):
                hdr_cells[i].text = col
            
            # 添加数据行
            for _, row in sorted_df.head(5).iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        else:
            doc.add_heading('2. 行业消费记录', level=1)
            add_text_to_doc("没有行业消费记录")

        doc.add_page_break()

        # 定义区域关键词库
        location_keywords = ["潢川县", "深圳市", "郑州市", "温县"]  # 可根据需要添加更多地区

        # 尝试可能的列名变体
        column_variants = ['交易对方', '对方', '收款方', '商户名称', 'counterpart', 'recipient']
        target_column = next((col for col in column_variants if col in df.columns), None)

        if target_column is None:
            available_cols = ", ".join(df.columns.tolist())
            raise KeyError(f"找不到交易对方列，可用的列有：{available_cols}")

        # 筛选包含地理关键词的交易并创建副本
        filtered_df = df[df[target_column].str.contains('|'.join(location_keywords), na=False, case=False, regex=True)].copy()

        # 检查金额列名称
        amount_variants = ['金额(元)', '金额', 'money', 'amount', '交易金额']
        amount_column = next((col for col in amount_variants if col in df.columns), None)

        if amount_column is None:
            available_cols = ", ".join(df.columns.tolist())
            raise KeyError(f"找不到金额列，可用的列有：{available_cols}")

        # 按金额倒序排列
        sorted_df = filtered_df.dropna(subset=[amount_column]).sort_values(by=amount_column, ascending=False)

        # 设置显示选项
        pd.set_option('display.max_rows', None)
        pd.set_option('display.unicode.ambiguous_as_wide', True)
        pd.set_option('display.unicode.east_asian_width', True)

        # 提取区域信息并统计前5个区域
        def extract_location(text, keywords):
            for keyword in keywords:
                if keyword in str(text):
                    return keyword
            return "其他"

        sorted_df['区域'] = sorted_df[target_column].apply(lambda x: extract_location(x, location_keywords))

        # 按区域分组并计算总金额
        region_stats = sorted_df.groupby('区域')[amount_column].sum().sort_values(ascending=False).head(5)

        # 打印结果
        doc.add_heading('3. 区域消费金额统计', level=1)
        add_text_to_doc("区域消费金额统计：", font_size=12, bold=True)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '区域'
        hdr_cells[1].text = '金额(元)'
        
        # 添加数据行
        for region, amount in region_stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = region
            row_cells[1].text = f"{amount:.2f}"

        # 绘制区域消费金额统计柱状图
        plt.figure(figsize=(10, 6))
        plt.bar(region_stats.index, region_stats.values)
        plt.xlabel('区域')
        plt.ylabel('金额(元)')
        plt.title('区域消费金额统计')
        chart_path = os.path.join(app.config['CHART_FOLDER'], 'region_chart.png')
        plt.savefig(chart_path)
        plt.close()
        doc.add_picture(chart_path, width=Inches(6))

        doc.add_page_break()

        # 按交易对方分组，计算交易金额总和，并筛选出交易金额前5的交易对方
        # 确保在进行数值计算前金额列已经转换为数值类型
        top_5_by_amount = df.groupby('交易对方')['金额(元)'].sum().abs().sort_values(ascending=False).head(5)

        # 按交易对方分组，计算交易数量，并筛选出交易数量前5的交易对方
        top_5_by_count = df.groupby('交易对方')['金额(元)'].count().sort_values(ascending=False).head(5)

        doc.add_heading('4. 交易金额前5的交易对方', level=1)
        add_text_to_doc('交易金额前5的交易对方：', font_size=12, bold=True)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '交易对方'
        hdr_cells[1].text = '金额(元)'
        
        # 添加数据行
        for entity, amount in top_5_by_amount.items():
            row_cells = table.add_row().cells
            row_cells[0].text = entity
            row_cells[1].text = f"{amount:.2f}"

        # 绘制交易金额前5的交易对方柱状图
        plt.figure(figsize=(10, 6))
        plt.bar(top_5_by_amount.index, top_5_by_amount.values)
        plt.xlabel('交易对方')
        plt.ylabel('金额(元)')
        plt.title('交易金额前5的交易对方')
        chart_path = os.path.join(app.config['CHART_FOLDER'], 'top_5_amount_chart.png')
        plt.savefig(chart_path)
        plt.close()
        doc.add_picture(chart_path, width=Inches(6))

        doc.add_page_break()

        doc.add_heading('5. 交易数量前5的交易对方', level=1)
        add_text_to_doc('交易数量前5的交易对方：', font_size=12, bold=True)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '交易对方'
        hdr_cells[1].text = '交易数量'
        
        # 添加数据行
        for entity, count in top_5_by_count.items():
            row_cells = table.add_row().cells
            row_cells[0].text = entity
            row_cells[1].text = str(count)

        # 绘制交易数量前5的交易对方柱状图
        plt.figure(figsize=(10, 6))
        plt.bar(top_5_by_count.index, top_5_by_count.values)
        plt.xlabel('交易对方')
        plt.ylabel('交易数量')
        plt.title('交易数量前5的交易对方')
        chart_path = os.path.join(app.config['CHART_FOLDER'], 'top_5_count_chart.png')
        plt.savefig(chart_path)
        plt.close()
        doc.add_picture(chart_path, width=Inches(6))

        # 保存Word文档
        doc.save(output_doc_path)

        # 验证文件是否成功保存
        if not os.path.exists(output_doc_path):
            raise Exception(f"DOCX保存失败，文件不存在: {output_doc_path}")

    except Exception as e:
        # 错误发生时尝试删除不完整的文件
        if os.path.exists(output_doc_path):
            try:
                os.remove(output_doc_path)
            except:
                pass
        raise  # 重新抛出异常，让上层处理

if __name__ == '__main__':
    app.run(debug=True)